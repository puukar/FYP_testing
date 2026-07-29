# DOCX-specific parser to handle tables and formatted text
from docx import Document
from pathlib import Path
from typing import Optional
import re


def extract_text_from_docx(docx_path: str) -> Optional[str]:
  
    try:
        if not Path(docx_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
            
        doc = Document(docx_path)
        text_parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
   
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
       
        combined = '\n'.join(text_parts)
        return combined if combined.strip() else None
        
    except Exception as e:
        print(f"[python-docx] Extraction error: {e}")
        return None